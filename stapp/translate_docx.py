import os
from dotenv import load_dotenv
import requests
import time
from docx import Document
import datetime
from glossary_database import get_glossary

load_dotenv()

X_RAPIDAPI_KEY = os.getenv("X-RapidAPI-Key")
X_RAPIDAPI_HOST = os.getenv("X-RapidAPI-Host")

# glossary = {
#     "적합기준": "Criteria",
#     "허용기준치": "Criteria",
#     "경고수준": "Alert Level",
#     "정확성": "Accuracy",
#     "허용범위": "Action level",
#     "조치수준": "Action level",
#     "원료의약품":"Active pharmaceutical Ingredient (API)",
#     "실 생산량":"Actual yield",
#     "전실":"Air-Lock",
#     "시험방법":"Analytical Procedure",
#     "무균 충전":"Aseptic filling",
#     "무균 기술 및 조작법":"Aseptic techniques and manipulations",
# }

def translate_text(text):
    if not text.strip():
        return text
    #Apply glossary corrections before making API call
    corrected_text = text
    glossary = get_glossary()  # retrieve the glossary from the database
    for ko_word, en_word in glossary.items():
        if ko_word in corrected_text:
            print(f"Replacing {ko_word} with {en_word}")
            corrected_text = corrected_text.replace(ko_word, en_word)

    url = "https://deepl-translator.p.rapidapi.com/translate"
    headers = {
        "content-type": "application/json",
        "X-RapidAPI-Key": X_RAPIDAPI_KEY,
        "X-RapidAPI-Host": X_RAPIDAPI_HOST
    }
    data = {
        "text": corrected_text,
        "source": "ko",
        "target": "en"
    }
    time.sleep(1)  # add delay between each API call
    response = requests.post(url, json=data, headers=headers)

    if response.status_code == 200:
        translation = response.json()['text']
        print(f"Translated: {text} to {translation}")
        return translation

    else:
        print(f"Translation failed for: {text}, status code: {response.status_code}, response: {response.text}")
        return text  # return original text if translation fails

def translate_docx(filename):
    start_time = time.time()
    doc = Document(filename)

    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        translated_text = translate_text(original_text)
        paragraph.add_run('\n'+translated_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    translated_text = translate_text(original_text)
                    paragraph.text = ''  # clear the original text
                    paragraph.add_run(translated_text)  # add the translated text

    output_filename = datetime.date.today().strftime('%Y%m%d') + 'translated_' + os.path.basename(filename)
    doc.save(output_filename)

    end_time = time.time()  # record the end time
    execution_time = end_time - start_time  # calculate the execution time
    print(f'Translation and file generation completed in {execution_time/60} minutes.')
    return output_filename




